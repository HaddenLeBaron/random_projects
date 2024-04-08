import os.path
import requests
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import base64
import docx

SCOPES = ["https://www.googleapis.com/auth/gmail.readonly"]
doc = docx.Document()
email_list = []
page_token = None

headers = False
target_sender = "Charlotte Colsch <charlotte.colsch@missionary.org>"
target_label = 'Label_5902939514796433665'


def parse_mail(dictionary):
    email_dict = dictionary

    formatted_date = email_dict['Date']

    paragraph = doc.add_paragraph('')
    paragraph.add_run(formatted_date + "\n" + email_dict["subject"]).bold = True
    p2 = doc.add_paragraph(email_dict['body'])


def main():
    """Shows basic usage of the Gmail API.
    Lists the user's Gmail labels.
    """
    creds = None
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open("token.json", "w") as token:
            token.write(creds.to_json())
    try:
        service = build("gmail", "v1", credentials=creds)
        results = service.users().labels().list(userId="me").execute()
        labels = results.get("labels", [])

        if not labels:
            print("No labels found.")
            return
        print("Labels:")
        for label in labels:
            print(f'name:{label["name"]}\nid:{label["id"]}')
    except HttpError as error:
        # TODO(developer) - Handle errors from gmail API.
        print(f"An error occurred: {error}")


def compile_label():
    global page_token
    done = False
    creds = None
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file("credentials.json", SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open("token.json", "w") as token:
            token.write(creds.to_json())
    # the Audrey tab Replace with appropriate Label ID when known
    to_compile = target_label
    try:
        # Call the Gmail API
        service = build("gmail", "v1", credentials=creds)
        results = service.users().messages().list(userId="me", labelIds=[to_compile], pageToken=page_token).execute()
        try:
            page_token = results["nextPageToken"]
        except KeyError:
            done = True
            print("Last page?")
        messages = results.get("messages", [])
        for message in messages:
            email_dict = {}
            msg = service.users().messages().get(userId="me", id=message['id']).execute()
            email_data = msg["payload"]
            is_response = False
            # TODO TURN ON FOR NEW PERSON
            print(email_data["headers"])
            for d in email_data["headers"]:
                if d["name"] == 'Date':
                    email_dict['Date'] = str(d["value"][0:16])
                if d["name"] == "Subject":
                    if "Re:" in d["value"]:
                        is_response = True
                    email_dict["subject"] = str(d["value"]).replace("\'", "'").replace("Fw:", "")
                # TODO Update From for new person
                if d["name"] == "From":
                    if d["value"] != target_sender:
                        is_response = True

            email_dict['body'] = ''
            try:
                for part in email_data['parts'][0:1]:
                    byte_code = base64.urlsafe_b64decode(part["body"]['data'])
                    text = byte_code.decode("utf-8")
                    formatted_string = str(text).split("Subject:")[-1].replace('\r\n', "\n").\
                        replace("&#39;", "'").replace("\'", "'")
                    email_dict["body"] += formatted_string
            except KeyError:
                try:
                    for part in email_data['parts'][0:1]:
                        for p in part['parts']:
                            byte_code = base64.urlsafe_b64decode(p["body"]['data'])
                            text = byte_code.decode("utf-8")
                            formatted_text = str(str(text).split('<div')[0])\
                                .replace('\r\n', " ").replace("&#39;", "'").replace("\'", "'")
                            email_dict["body"] += formatted_text
                except KeyError:
                    print("still didn't work dang it.")
            if not is_response:
                email_list.append(email_dict)
    except HttpError as error:
        # TODO(developer) - Handle errors from gmail API.
        print(f"An error occurred: {error}")
    return done


if __name__ == "__main__":
    if headers:
        main()    # for finding label IDs
    compile_label()
    while page_token:
        if compile_label():
            break
        print("Next page!")
    for email in email_list[::-1]:
        parse_mail(email)
    doc.save("document.docx")
